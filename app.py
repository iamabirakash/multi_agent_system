from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from typing import Any
import json
import threading
import uuid
import asyncio
import re
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from docx import Document

from pipeline import run_research_pipeline


class ResearchRequest(BaseModel):
    topic: str = Field(min_length=3, max_length=200)


app = FastAPI(title="Multi Agent Research API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
def health_check() -> dict:
    return {"status": "ok"}


def _to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        parts = []
        for item in value:
            parts.append(_to_text(item))
        return "\n".join(part for part in parts if part).strip()
    if isinstance(value, dict):
        if "text" in value and isinstance(value["text"], str):
            return value["text"]
        if "content" in value:
            return _to_text(value["content"])
        try:
            return json.dumps(value, ensure_ascii=False, indent=2)
        except Exception:
            return str(value)
    return str(value)


def _slugify_filename(value: str) -> str:
    clean = re.sub(r"[^a-zA-Z0-9\s_-]", "", value or "").strip().lower()
    clean = re.sub(r"[\s_-]+", "-", clean)
    return clean[:80] or "research-report"


STEP_LABELS = {
    "search": "Searching trustworthy sources",
    "reader": "Scraping most relevant article",
    "writer": "Writing structured research report",
    "critic": "Generating critic feedback",
}

JOBS = {}
JOBS_LOCK = threading.Lock()


def _initial_steps() -> list[dict]:
    return [
        {"key": "search", "label": STEP_LABELS["search"], "status": "pending"},
        {"key": "reader", "label": STEP_LABELS["reader"], "status": "pending"},
        {"key": "writer", "label": STEP_LABELS["writer"], "status": "pending"},
        {"key": "critic", "label": STEP_LABELS["critic"], "status": "pending"},
    ]


def _set_step_status(job_id: str, step_key: str, status: str) -> None:
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            return
        for step in job["steps"]:
            if step["key"] == step_key:
                step["status"] = status
                break


def _run_job(job_id: str, topic: str) -> None:
    def progress_callback(step_key: str, status: str) -> None:
        _set_step_status(job_id, step_key, status)

    try:
        result = run_research_pipeline(topic, verbose=False, progress_callback=progress_callback)
        with JOBS_LOCK:
            if job_id in JOBS:
                JOBS[job_id]["status"] = "done"
                JOBS[job_id]["report"] = _to_text(result.get("report", ""))
                for step in JOBS[job_id]["steps"]:
                    step["status"] = "done"
    except Exception as exc:
        with JOBS_LOCK:
            if job_id in JOBS:
                JOBS[job_id]["status"] = "error"
                JOBS[job_id]["error"] = f"Pipeline failed: {exc}"


@app.post("/api/research/start")
def start_research(payload: ResearchRequest) -> dict:
    job_id = str(uuid.uuid4())
    with JOBS_LOCK:
        JOBS[job_id] = {
            "job_id": job_id,
            "topic": payload.topic,
            "status": "running",
            "steps": _initial_steps(),
            "report": "",
            "error": "",
        }
    thread = threading.Thread(target=_run_job, args=(job_id, payload.topic), daemon=True)
    thread.start()
    return {"job_id": job_id}


@app.get("/api/research/{job_id}")
def get_research_status(job_id: str) -> dict:
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        return {
            "job_id": job["job_id"],
            "topic": job["topic"],
            "status": job["status"],
            "steps": job["steps"],
            "report": job["report"],
            "error": job["error"],
        }


@app.get("/api/research/{job_id}/stream")
async def stream_research_status(job_id: str):
    with JOBS_LOCK:
        if job_id not in JOBS:
            raise HTTPException(status_code=404, detail="Job not found.")

    async def event_generator():
        last_payload = ""

        while True:
            with JOBS_LOCK:
                job = JOBS.get(job_id)
                if not job:
                    payload = {
                        "job_id": job_id,
                        "status": "error",
                        "error": "Job not found.",
                    }
                    is_terminal = True
                else:
                    payload = {
                        "job_id": job["job_id"],
                        "topic": job["topic"],
                        "status": job["status"],
                        "steps": job["steps"],
                        "report": job["report"],
                        "error": job["error"],
                    }
                    is_terminal = job["status"] in ("done", "error")

            payload_text = json.dumps(payload, ensure_ascii=False)
            if payload_text != last_payload:
                yield f"event: update\ndata: {payload_text}\n\n"
                last_payload = payload_text

            if is_terminal:
                yield "event: end\ndata: {}\n\n"
                break

            await asyncio.sleep(0.5)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/research/{job_id}/export/{fmt}")
def export_research_report(job_id: str, fmt: str):
    allowed_formats = {"txt", "md", "pdf", "docx"}
    if fmt not in allowed_formats:
        raise HTTPException(status_code=400, detail=f"Unsupported format. Use one of: {', '.join(sorted(allowed_formats))}")

    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        if job.get("status") != "done":
            raise HTTPException(status_code=400, detail="Research job is not complete yet.")
        report = _to_text(job.get("report", ""))
        topic = _to_text(job.get("topic", ""))

    if not report.strip():
        raise HTTPException(status_code=400, detail="No report available to export.")

    filename_base = _slugify_filename(topic)

    if fmt == "txt":
        content = report
        media_type = "text/plain; charset=utf-8"
    elif fmt == "md":
        content = report
        media_type = "text/markdown; charset=utf-8"
    elif fmt == "pdf":
        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        margin = 18 * mm
        y = height - margin

        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(margin, y, topic or "Research Report")
        y -= 12 * mm

        pdf.setFont("Helvetica", 10)
        max_chars = 105
        for paragraph in report.splitlines():
            line = paragraph.strip()
            if not line:
                y -= 5 * mm
                if y < margin:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    y = height - margin
                continue
            chunks = [line[i:i + max_chars] for i in range(0, len(line), max_chars)]
            for chunk in chunks:
                pdf.drawString(margin, y, chunk)
                y -= 5 * mm
                if y < margin:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    y = height - margin

        pdf.save()
        content = buffer.getvalue()
        media_type = "application/pdf"
    else:
        document = Document()
        document.add_heading(topic or "Research Report", 0)
        for paragraph in report.splitlines():
            document.add_paragraph(paragraph)
        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return Response(
        content=content,
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename_base}.{fmt}"
        },
    )
